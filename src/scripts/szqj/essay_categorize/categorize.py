import os

import docx
import jieba

from utils.text_utilizer import replace_punctuation
from szqj.essay_categorize.dt import *


def create_file_tree():
    file_tree = {}
    root_path = "../../../../data/nlp/文章标签库/"

    lvl1_list = os.listdir(root_path)

    count = 0
    success = 0

    for item1 in lvl1_list:
        lvl1_tree = {}
        for item2 in os.listdir(root_path + item1):

            f_list = []
            for file in os.listdir("{}{}/{}/".format(root_path, item1, item2)):
                doc = docx.Document("{}{}/{}/{}".format(root_path, item1, item2, file))
                local_success = 0
                for par in doc.paragraphs[:3]:
                    par = par.text
                    if "标签：" in par.replace(":", "：") or "关键词：" in par.replace(":", "："):
                        # print(par)
                        local_success += 1
                        success += 1
                if local_success == 0:
                    print(doc.paragraphs[0].text)
                    print(item1, item2, file)
                count += 1
                content = "\n".join(txt.text for txt in doc.paragraphs)
                f_list.append(content)
            lvl1_tree[item2] = f_list
        file_tree[item1] = lvl1_tree

    with open("../../../../data/output/file_tree.pickle", "wb") as f:
        pickle.dump(file_tree, f)


def reconstruct():
    with open("../../../../data/output/file_tree.pickle", "rb") as f:
        file_tree = pickle.load(f)

    success = 0
    tot = 0

    new_tree = {}

    for lvl1 in file_tree:
        lvl1_tree = {}
        for lvl2 in file_tree[lvl1]:
            new_f_list = []
            f_list = file_tree[lvl1][lvl2]
            for f in f_list:
                paragraphs = f.split("\n")
                local_success = 0
                for par in paragraphs[:3]:
                    if "标签：" in par.replace(":", "：") or "关键词：" in par.replace(":", "："):
                        local_success += 1
                        success += 1
                        paragraphs.remove(par)
                        new_f_list.append((par, "\n".join(paras for paras in paragraphs)))
                        break
                if local_success == 0:
                    print(paragraphs[0])
                tot += 1
            lvl1_tree[lvl2] = new_f_list
        new_tree[lvl1] = lvl1_tree

    with open("../../../../data/output/file_tree.pickle", "wb") as f:
        pickle.dump(new_tree, f)


def gen_onehot():
    with open("../../../../data/output/file_tree.pickle", "rb") as f:
        file_tree = pickle.load(f)
    with open("../../../../data/nlp/stop_words.pickle", "rb") as stopfile:
        stopwords = pickle.load(stopfile)
    with open("../../../../data/nlp/idf.pickle", "rb") as file:
        idf = pickle.load(file)
    kws = []
    for lvl1 in file_tree:
        for lvl2 in file_tree[lvl1]:
            for item in file_tree[lvl1][lvl2]:
                tag = item[0]
                content = item[1]
                tag = replace_punctuation(tag, replacement=" ", exception="-&+").replace("  ", " ").replace(" , ", " ")
                wl = tag.split(" ")
                if "" in wl:
                    wl.remove("")
                for w in wl:
                    try:
                        widf = idf[w]
                    except KeyError as e:
                        widf = 2 ** 10
                    if w in stopwords:
                        pass
                    # elif widf < 2500 or widf > 5000:
                    #     pass
                    else:
                        kws.append(w)

    with open("../../../../data/nlp/essay_onehot.pickle", "wb") as f:
        pickle.dump(list(set(kws)), f)


# gen_onehot()
# 
# with open("../../../../data/output/file_tree.pickle", "rb") as f:
#     file_tree = pickle.load(f)
# 
# with open("../../../../data/nlp/essay_onehot.pickle", "rb") as f:
#     kws = pickle.load(f)
# print(len(kws))


if __name__ == "__main__":
    _ = 1

    dic_path = "../../../../data/output/account_name_unique_jieba.txt"
    jieba.load_userdict(dic_path)

    counter = 0
    with open("../../../../data/output/file_tree.pickle", "rb") as f:
        file_tree = pickle.load(f)

    with open("../../../../data/nlp/essay_onehot__.pickle", "rb") as f:
        kws = pickle.load(f)

    with open("../../../../data/nlp/idf_.pickle", "rb") as f:
        idf = pickle.load(f)

    # input data: the last col is label
    training = []
    for lvl1 in file_tree:
        for lvl2 in file_tree[lvl1]:
            for item in file_tree[lvl1][lvl2][:-3]:
                tag = item[0]
                content = item[1]
                content = " ".join(jieba.cut(content)).replace("\n", "").replace("  ", " ").split(" ")

                onehot = [0] * len(kws)
                for i in range(len(kws)):
                    onehot[i] += content.count(kws[i]) / idf[kws[i]]

                if onehot == [0] * len(kws):
                    # counter += 1
                    onehot.append(1)
                else:
                    onehot.append(0)
                onehot.append("{}-{}".format(lvl1, lvl2))
                training.append(onehot)

    # testing data
    testing = []
    testing_essay = []
    for lvl1 in file_tree:
        for lvl2 in file_tree[lvl1]:
            for item in file_tree[lvl1][lvl2][-3:]:
                tag = item[0]
                content = item[1]
                content = " ".join(jieba.cut(content)).replace("\n", "").replace("  ", " ").split(" ")

                onehot = [0] * len(kws)
                for i in range(len(kws)):
                    onehot[i] += content.count(kws[i]) / idf[kws[i]]
                if onehot == [0] * len(kws):
                    counter += 1
                    onehot.append(1)
                else:
                    onehot.append(0)
                onehot.append("{}-{}".format(lvl1, lvl2))
                testing.append(onehot)
                testing_essay.append((lvl1, lvl2, content))

    # Column labels.
    # These are used only to print the tree.
    header = []
    for _ in range(len(testing[0])):
        header.append('{}'.format(_))
    my_tree = build_tree(training)

    # """测试"""
    # 正确计数

    # 将训练结果保存以防重复训练便于使用
    with open('../../../../data/output/category_dt.pickle', 'wb') as f:
        pickle.dump(my_tree, f, protocol=pickle.HIGHEST_PROTOCOL)
    corr = 0
    with open('../../../../data/output/category_dt.pickle', 'rb') as f:
        my_tree = pickle.load(f)
    # print_tree(my_tree)
    for i in range(len(testing)):
        row = testing[i]
        act = row[-1]
        possibilities = print_leaf(classify(row, my_tree))
        temp = 0
        for case in possibilities.keys():
            # print(float(possibilities[case][:-1]))
            if float(possibilities[case][:-1]) > temp:
                prediction = case
                temp = float(possibilities[case][:-1])
        # prediction = list(classify(row, my_tree).keys())[0]
        # print(type(act))
        if prediction == act:
            corr += 1
        else:
            print(
                "Actual: {} === Predicted: {} === Prob: {}".format(act, prediction, print_leaf(classify(row, my_tree))))
            # print(str(testing_essay[i]) + "\n")
        # print("Actual: {}. Predicted: {}".format(act, prediction))
    print('Accuracy: {}%'.format(100 * (corr / len(testing))))
